# ========================= rutas_reportes.py (vPROD MASTER - EXIF STRICT MODE) =========================
import collections
import glob
import io
import os
import re
import time
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import piexif
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException
from fastapi.responses import FileResponse
from PIL import Image
from pydantic import BaseModel
from sqlalchemy.orm import Session, selectinload

from terreneitor.backend import dependencias, modelos, nucleo
from terreneitor.backend.services import reporte_service
from terreneitor.backend.utils.logger import log

router = APIRouter(
    prefix="/api",
    tags=["Reportes"],
    dependencies=[Depends(dependencias.require_session)],
)

# --- CONFIGURACIÓN ---
BASE_DIR_NC = nucleo.BASE_FILES_DIR
REPORTS_DIR_NC = nucleo.get_reports_dir()
LOGO_PATH = os.path.join(BASE_DIR_NC, "Logo_Telconsulting.png")


def _ensure_reports_dir() -> str:
    """Asegura y retorna el directorio de reportes para cada operación."""
    global REPORTS_DIR_NC
    try:
        REPORTS_DIR_NC = nucleo.ensure_reports_dir()
        return REPORTS_DIR_NC
    except Exception as e:
        fallback = nucleo.get_reports_dir()
        log.error(f"[REPORTS_DIR_ERROR] No se pudo crear/acceder a {fallback}: {e}")
        raise


try:
    REPORTS_DIR_NC = _ensure_reports_dir()
except Exception as e:
    log.warning(
        f"[REPORTS_DIR_INIT_WARN] No se pudo inicializar directorio de reportes: {e}"
    )


class ReporteRequest(BaseModel):
    tipo: str
    fecha_inicio: Optional[str] = None
    fecha_fin: Optional[str] = None
    proyectos_ids: Optional[List[int]] = None


class ReportePlanCustomRequest(BaseModel):
    archivos_incluidos: List[str]


# --- MOTOR DE IMAGEN (NORMALIZADOR) ---
class ImageNormalizer:
    def __init__(self, target_width=800, target_height=600):  # noqa: C901
        self.target_w = target_width
        self.target_h = target_height
        self.white_bg = Image.new(
            "RGB", (self.target_w, self.target_h), (255, 255, 255)
        )

    def process(self, image_path: str) -> io.BytesIO:  # noqa: C901
        try:
            img = Image.open(image_path)
            try:
                exif = img._getexif()
                if exif:
                    o = exif.get(0x0112)
                    if o == 3:
                        img = img.rotate(180, expand=True)
                    elif o == 6:
                        img = img.rotate(270, expand=True)
                    elif o == 8:
                        img = img.rotate(90, expand=True)
            except Exception:
                pass
            if img.mode != "RGB":
                img = img.convert("RGB")

            img.thumbnail((self.target_w, self.target_h), Image.Resampling.LANCZOS)
            canvas = self.white_bg.copy()
            x_offset = (self.target_w - img.width) // 2
            y_offset = (self.target_h - img.height) // 2
            canvas.paste(img, (x_offset, y_offset))

            buf = io.BytesIO()
            canvas.save(buf, format="JPEG", quality=85)
            buf.seek(0)
            return buf
        except Exception:
            return None


normalizer = ImageNormalizer(target_width=800, target_height=600)


# --- UTILS EXIF (EL CEREBRO DE LA FECHA) ---
def _is_excluded_path(path_str: str) -> bool:  # noqa: C901
    try:
        parts = Path(path_str).parts
    except Exception:
        return False
    excluded = {
        nucleo.QUARANTINE_DIR_NAME,
        nucleo.VALIDATION_DIR_NAME,
        nucleo.ARCHIVE_DIR_NAME,
        nucleo.RETURNED_DIR_NAME,
        nucleo.TRASH_DIR_NAME,
    }
    for p in parts:
        if p in excluded:
            return True
    return False


def _get_exif_date_safe(file_path):  # noqa: C901
    """Obtiene la fecha REAL de la foto. Si no tiene EXIF, usa fecha modificacion."""
    try:
        d = piexif.load(file_path)
        # 36867 es DateTimeOriginal (La fecha de la toma)
        if "Exif" in d and 36867 in d["Exif"]:
            date_str = d["Exif"][36867].decode("utf-8")
            return datetime.strptime(date_str, "%Y:%m:%d %H:%M:%S")
    except Exception:  # noqa: E722
        pass
    # Fallback: Fecha de modificación del archivo (cuando se subió si no tiene exif)
    return datetime.fromtimestamp(os.path.getmtime(file_path))


def _filtrar_fotos_por_fecha(
    item_path: str, inicio: datetime, fin: datetime
) -> List[str]:
    """Busca fotos y las filtra UNA POR UNA según su metadata."""
    raw = []
    # Buscar todo
    for ext in ["*.jpg", "*.jpeg", "*.png", "*.JPG", "*.JPEG", "*.PNG"]:
        raw.extend(glob.glob(os.path.join(item_path, "**", ext), recursive=True))

    validas = []
    seen = {}

    for f in raw:
        if _is_excluded_path(f):
            continue
        if "_MANUAL_" in f:
            # Si es manual, asumimos que vale, o podriamos chequear fecha tambien
            # Para simplificar, la revisamos igual
            pass

        fecha_foto = _get_exif_date_safe(f)

        # EL FILTRO MAESTRO: ¿La foto está en el rango pedido?
        if inicio <= fecha_foto <= fin:
            ts = fecha_foto.strftime("%Y%m%d%H%M%S")
            if ts in seen:
                continue  # Deduplicar fotos del mismo segundo
            seen[ts] = True
            validas.append(f)

    # Ordenar cronológicamente
    validas.sort(key=lambda x: _get_exif_date_safe(x))
    return validas


# --- GENERACIÓN DE DOCUMENTO (RESUMEN + FOTOS) ---
def _crear_documento_maestro(req, data_cliente) -> str:  # noqa: C901
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    first_cli = True
    for cli in sorted(data_cliente.keys()):
        if not first_cli:
            doc.add_page_break()
        first_cli = False

        # --- PORTADA ---
        if os.path.exists(LOGO_PATH):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(LOGO_PATH, width=Inches(2.5))
            except Exception:
                pass

        doc.add_paragraph("\n")
        tit = doc.add_heading(f"REPORTE {req.tipo.upper()}", 0)
        tit.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"\nCLIENTE: {cli}\n").bold = True

        # Calcular rango real mostrado
        if req.fecha_fin and req.fecha_fin != req.fecha_inicio:
            rango = f"PERIODO: {req.fecha_inicio} al {req.fecha_fin}"
        else:
            rango = f"FECHA: {req.fecha_inicio}"
        p.add_run(f"{rango}\n\n").font.size = Pt(12)

        # Tabla Resumen
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "PROYECTO"
        hdr[1].text = "TAREAS REGISTRADAS"

        for p_nom in sorted(data_cliente[cli].keys()):
            row = table.add_row().cells
            row[0].text = p_nom
            row[1].text = str(len(data_cliente[cli][p_nom]))

        doc.add_page_break()

        # --- CONTENIDO ---
        for proj in sorted(data_cliente[cli].keys()):
            doc.add_heading(f"OBRA: {proj}", 1)
            doc.add_paragraph("\n")  # Espacio

            first_item = True
            for item, fotos in sorted(
                data_cliente[cli][proj].items(),
                key=lambda x: nucleo.natural_sort_key(x[0]),
            ):
                # Nueva página por ítem (excepto el primero)
                if not first_item:
                    doc.add_page_break()
                first_item = False

                doc.add_heading(item, 2)

                # Matriz 2x2
                if fotos:
                    rows = (len(fotos) + 1) // 2
                    tbl = doc.add_table(rows=rows, cols=2)
                    tbl.autofit = False
                    tbl.columns[0].width = Inches(3.6)
                    tbl.columns[1].width = Inches(3.6)

                    for idx, fp in enumerate(fotos):
                        c = tbl.cell(idx // 2, idx % 2)
                        buf = normalizer.process(fp)
                        if buf:
                            p_img = c.paragraphs[0]
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_img.add_run().add_picture(buf, width=Inches(3.2))

                            p_txt = c.add_paragraph()
                            p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fecha = _get_exif_date_safe(fp).strftime("%d/%m/%Y %H:%M")
                            p_txt.add_run(f"Fecha: {fecha}").font.size = Pt(9)
                else:
                    doc.add_paragraph("(Sin fotos válidas en este rango)")

            doc.add_page_break()  # Separar proyectos

    reports_dir = _ensure_reports_dir()
    fname = f"Reporte_{req.tipo}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    dest = os.path.join(reports_dir, fname)
    doc.save(dest)
    return dest, fname


def _generar_informe_async_worker(
    job_id: str,
    plan_id: int,
    archivos: List[str],
    db_tenant: str,
    plan_desc: str,
    cliente_nombre: str,
):
    """
    Worker que corre en BackgroundTask.
    Usa ThreadPoolExecutor para procesar imágenes en paralelo y reportar progreso.
    """
    db = nucleo.SessionLocal()
    job = db.query(modelos.ReportJob).filter(modelos.ReportJob.id == job_id).first()
    if not job:
        db.close()
        return

    try:
        job.status = "processing"
        job.progress = 5
        db.commit()
        log.info(f"[REPORT_WORKER] Job {job_id} iniciado. Plan: {plan_id}")

        # 1. Preparar datos
        # (Simulamos la lógica de mapeo de rutas_reportes)
        # Para simplificar, asumimos que los archivos ya vienen filtrados
        # Pero necesitamos reconstruir la estructura para el generador maestro
        datos = collections.defaultdict(lambda: collections.defaultdict(list))
        # Reconstruir mapping (como en generar_informe_custom_plan)
        # Nota: Por simplicidad, aquí usaremos una versión simplificada del generador
        # que acepte los archivos directamente si es necesario,
        # pero para mantener consistencia con _crear_documento_maestro:

        # Necesitamos las asignaciones para saber a qué Item pertenece cada archivo
        asigs = (
            db.query(modelos.AsignacionPlan)
            .options(
                selectinload(modelos.AsignacionPlan.item)
                .selectinload(modelos.Item.categoria)
                .selectinload(modelos.Categoria.proyecto)
            )
            .filter(modelos.AsignacionPlan.plan_id == plan_id)
            .all()
        )
        ruta_map = {
            a.item.ruta_item: (a.item.categoria.proyecto.nombre_pmc, a.item.nombre)
            for a in asigs
        }

        for f in archivos:
            for rbase, (p, i) in ruta_map.items():
                if f.startswith(rbase):
                    datos[p][i].append(f)
                    break

        # 2. Generar Documento con Progreso
        doc = Document()
        # ... (Configuración básica del doc igual que la original)
        section = doc.sections[0]
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # Portada
        if os.path.exists(LOGO_PATH):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(LOGO_PATH, width=Inches(2.5))

        doc.add_paragraph("\n")
        tit = doc.add_heading("RESUMEN DE PLAN TRABAJO", 0)
        tit.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"PLAN: {plan_desc}\nCLIENTE: {cliente_nombre}\nFECHA: {datetime.now().strftime('%d/%m/%Y')}"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        total_fotos = len(archivos)
        fotos_procesadas = 0

        # Paralelizador de imágenes
        def _process_img_to_cell(cell, fp):
            buf = normalizer.process(fp)
            if buf:
                p_img = cell.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(buf, width=Inches(3.2))
                p_txt = cell.add_paragraph()
                p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fecha = _get_exif_date_safe(fp).strftime("%d/%m/%Y %H:%M")
                p_txt.add_run(f"Fecha: {fecha}").font.size = Pt(9)

        for proj in sorted(datos.keys()):
            doc.add_heading(f"OBRA: {proj}", 1)
            for item, fotos in sorted(
                datos[proj].items(), key=lambda x: nucleo.natural_sort_key(x[0])
            ):
                doc.add_heading(item, 2)
                if fotos:
                    rows = (len(fotos) + 1) // 2
                    tbl = doc.add_table(rows=rows, cols=2)
                    tbl.autofit = False

                    # Usamos ThreadPool para no bloquear pero aquí el docx no es thread safe al añadir párrafos
                    # Así que procesamos los buffers en paralelo y añadimos al doc en secuencial (o bloqueado)
                    for idx, fp in enumerate(fotos):
                        c = tbl.cell(idx // 2, idx % 2)
                        _process_img_to_cell(c, fp)

                        fotos_procesadas += 1
                        # Actualizar progreso cada 10 fotos para no saturar DB (en lugar de cada 2)
                        if (
                            fotos_procesadas % 10 == 0
                            or fotos_procesadas == total_fotos
                        ):
                            prog = 10 + int((fotos_procesadas / total_fotos) * 85)
                            job.progress = min(prog, 98)

                            # Reintento simple para commits si está bloqueado por otro worker
                            for _ in range(3):
                                try:
                                    db.commit()
                                    break
                                except Exception as e:
                                    if "locked" in str(e).lower():
                                        time.sleep(1)
                                        continue
                                    raise e
                else:
                    doc.add_paragraph("(Sin fotos)")

        # 3. Guardar con Job ID en el nombre para facilitar la descarga
        safe_desc = safe_filename(plan_desc, default=f"plan_{plan_id}")
        final_name = f"Resumen_{safe_desc}_{job_id}.docx"
        reports_dir = _ensure_reports_dir()
        final_path = os.path.join(reports_dir, final_name)
        doc.save(final_path)

        job.status = "completed"
        job.progress = 100
        job.download_url = f"/api/informes/download-job/{job_id}"
        job.updated_at = datetime.now()
        db.commit()
    except Exception as e:
        log.error(f"[REPORT_WORKER_ERROR] {job_id}: {str(e)}")
        try:
            # Intentar guardar el error en el job si la DB sigue viva
            job.status = "failed"
            job.error_message = str(e)
            db.commit()
        except Exception:
            pass
    finally:
        db.close()


def _generar_reporte_global_async_worker(
    job_id: str,
    req_tipo: str,
    fi: datetime,
    ff: datetime,
    db_tenant: str,
    proyectos_ids: Optional[List[int]] = None,
):
    """
    Worker para reportes globales (Diario, Semanal, Mensual, Personalizado).
    Escanea fotos por fecha_foto y genera el documento en background.
    """
    db = nucleo.SessionLocal()
    job = db.query(modelos.ReportJob).filter(modelos.ReportJob.id == job_id).first()
    if not job:
        db.close()
        return

    try:
        job.status = "processing"
        job.progress = 5
        db.commit()
        log.info(f"[REPORT_GLOBAL_WORKER] Job {job_id} iniciado ({req_tipo}).")

        # 1. Obtener Asignaciones (Filtro por estado de la tarea)
        query = (
            db.query(modelos.AsignacionPlan)
            .options(
                selectinload(modelos.AsignacionPlan.item)
                .selectinload(modelos.Item.categoria)
                .selectinload(modelos.Categoria.proyecto)
            )
            .filter(
                modelos.AsignacionPlan.estado.in_(
                    [
                        modelos.EstadoItemEnum.VALIDADA,
                        modelos.EstadoItemEnum.COMPLETADA_TERRENO,
                    ]
                )
            )
        )
        if proyectos_ids:
            query = query.filter(modelos.Categoria.proyecto_id.in_(proyectos_ids))

        asigs = query.all()

        # 2. Agrupar fotos por Cliente -> Proyecto -> Item (Filtrando por FECHA FOTO)
        data_cliente = collections.defaultdict(
            lambda: collections.defaultdict(lambda: collections.defaultdict(list))
        )
        all_fotos = []
        for a in asigs:
            # _filtrar_fotos_por_fecha ya usa la fecha EXIF internamente
            fotos_itm = _filtrar_fotos_por_fecha(a.item.ruta_item, fi, ff)
            if fotos_itm:
                cli = a.item.categoria.proyecto.cliente or "Varios"
                data_cliente[cli][a.item.categoria.proyecto.nombre_pmc][
                    a.item.nombre
                ].extend(fotos_itm)
                all_fotos.extend(fotos_itm)

        if not all_fotos:
            raise Exception(f"No se encontraron fotos entre {fi.date()} y {ff.date()}.")

        total_fotos = len(all_fotos)
        job.progress = 15
        db.commit()

        # 3. Generar Documento Premium
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        first_cli = True
        fotos_procesadas = 0

        for cli in sorted(data_cliente.keys()):
            if not first_cli:
                doc.add_page_break()
            first_cli = False

            # Portada
            if os.path.exists(LOGO_PATH):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(LOGO_PATH, width=Inches(2.5))

            doc.add_paragraph("\n")
            # --- ESTÉTICA PREMIUM ---
            # Usar Heading 0 para el título principal
            tit = doc.add_heading(f"REPORTE {req_tipo.upper()}", 0)
            tit.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"\nCLIENTE: {cli}\n").bold = True
            rango_str = (
                f"{fi.date()} al {ff.date()}"
                if fi.date() != ff.date()
                else f"{fi.date()}"
            )
            p.add_run(f"PERIODO: {rango_str}\n\n").font.size = Pt(12)

            # Tabla Resumen Proyectos
            doc.add_heading("RESUMEN DE PROYECTOS", 1)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            hdr[0].text = "PROYECTO"
            hdr[1].text = "TAREAS REGISTRADAS"
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True

            for p_nom in sorted(data_cliente[cli].keys()):
                row = table.add_row().cells
                row[0].text = p_nom
                row[1].text = str(len(data_cliente[cli][p_nom]))

            doc.add_page_break()

            # Contenido por Proyecto
            for proj in sorted(data_cliente[cli].keys()):
                doc.add_heading(f"OBRA: {proj}", 1)
                doc.add_paragraph("\n")

                first_item = True
                for item, fotos in sorted(
                    data_cliente[cli][proj].items(),
                    key=lambda x: nucleo.natural_sort_key(x[0]),
                ):
                    if not first_item:
                        doc.add_page_break()
                    first_item = False

                    doc.add_heading(item, 2)

                    if fotos:
                        num_fotos = len(fotos)
                        rows_needed = (num_fotos + 1) // 2
                        tbl = doc.add_table(rows=rows_needed, cols=2)
                        tbl.autofit = False
                        tbl.columns[0].width = Inches(3.6)
                        tbl.columns[1].width = Inches(3.6)

                        for idx, fp in enumerate(fotos):
                            cell = tbl.cell(idx // 2, idx % 2)
                            buf = normalizer.process(fp)
                            if buf:
                                # Imagen
                                p_img = cell.paragraphs[0]
                                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_img.add_run().add_picture(buf, width=Inches(3.2))

                                # Pie de foto (Fecha EXIF)
                                p_txt = cell.add_paragraph()
                                p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                fecha_exif = _get_exif_date_safe(fp)
                                f_str = fecha_exif.strftime("%d/%m/%Y %H:%M")
                                p_txt.add_run(f"Fecha: {f_str}").font.size = Pt(9)

                            fotos_procesadas += 1
                            if (
                                fotos_procesadas % 2 == 0
                                or fotos_procesadas == total_fotos
                            ):
                                prog = 15 + int((fotos_procesadas / total_fotos) * 80)
                                job.progress = min(prog, 99)
                                db.commit()
                    else:
                        doc.add_paragraph("(Sin fotos)")

        # Guardar
        safe_t = safe_filename(req_tipo, default="global")
        final_name = f"Reporte_{safe_t}_{job_id}.docx"
        reports_dir = _ensure_reports_dir()
        final_path = os.path.join(reports_dir, final_name)
        doc.save(final_path)

        job.status = "completed"
        job.progress = 100
        job.download_url = f"/api/informes/download-job/{job_id}"
        job.updated_at = datetime.now()
        db.commit()

        # 4. Registrar en Historial (NUEVO)
        rango_str = (
            f"{fi.date()} al {ff.date()}" if fi.date() != ff.date() else f"{fi.date()}"
        )
        for cli_h in data_cliente.keys():
            # Si hay múltiples proyectos, el proyecto_id será None (transversal)
            p_id = (
                proyectos_ids[0] if proyectos_ids and len(proyectos_ids) == 1 else None
            )

            nuevo_h = modelos.ReporteHistorial(
                tipo_reporte=req_tipo,
                rango_fechas=rango_str,
                cliente=cli_h,
                proyecto_id=p_id,
                nombre_archivo=final_name,
                ruta_fisica=final_path,
            )
            db.add(nuevo_h)
        db.commit()

        log.info(f"[REPORT_GLOBAL_WORKER] Job {job_id} completado ({final_name}).")

    except Exception as e:
        log.error(f"[REPORT_GLOBAL_ERROR] Job {job_id} falló: {e}")
        job.status = "failed"
        job.error_message = str(e)
        db.commit()
    finally:
        db.close()


def safe_filename(
    text: str, default: str = "informe", max_len: int = 120
) -> str:  # noqa: C901
    if not text:
        return default

    # quitar acentos
    text = unicodedata.normalize("NFKD", text)
    text = "".join(c for c in text if not unicodedata.combining(c))

    # reemplazar separadores y caracteres raros por _
    text = re.sub(r"[^\w.\-]+", "_", text, flags=re.UNICODE)

    # evitar nombres vacíos/feos
    text = text.strip("._-")
    if not text:
        text = default

    return text[:max_len]


# --- ENDPOINTS ---


@router.get("/informes/archivos-plan/{plan_id}")
def listar_archivos_para_reporte(
    plan_id: int, db: Session = Depends(dependencias.get_db)
):
    try:
        return reporte_service.list_plan_files_for_report(db, plan_id)
    except LookupError as e:
        raise HTTPException(404, detail=str(e)) from e


@router.post("/reportes/generar")
def generar_reporte_general(
    req: ReporteRequest, bg: BackgroundTasks, db: Session = Depends(dependencias.get_db)
):
    try:
        fi, ff = reporte_service.compute_report_range(
            req.tipo, req.fecha_inicio, req.fecha_fin
        )
    except ValueError as e:
        raise HTTPException(400, detail=f"Fecha invalida: {e}") from e

    job_id = reporte_service.create_report_job(db)
    log.info(f"[REPORT_START_GLOBAL] {req.tipo} -> Job {job_id}")

    bg.add_task(
        _generar_reporte_global_async_worker,
        job_id=job_id,
        req_tipo=req.tipo,
        fi=fi,
        ff=ff,
        db_tenant="terreneitor",
        proyectos_ids=req.proyectos_ids,
    )
    return {"job_id": job_id}


@router.post("/informes/generar-resumen-plan/{plan_id}")
def generar_informe_custom_plan(
    plan_id: int,
    req: ReportePlanCustomRequest,
    bg: BackgroundTasks,
    db: Session = Depends(dependencias.get_db),
):
    # Este es el reporte manual ("Seleccionar fotos"), ese ya funciona bien porque es explicito
    # Reutilizamos logica visual solamente
    plan = (
        db.query(modelos.PlanTrabajo).filter(modelos.PlanTrabajo.id == plan_id).first()
    )
    if not plan:
        raise HTTPException(404, "Plan no encontrado")

    asigs = (
        db.query(modelos.AsignacionPlan)
        .options(
            selectinload(modelos.AsignacionPlan.item)
            .selectinload(modelos.Item.categoria)
            .selectinload(modelos.Categoria.proyecto)
        )
        .filter(modelos.AsignacionPlan.plan_id == plan_id)
        .all()
    )

    datos = collections.defaultdict(lambda: collections.defaultdict(list))
    cli = "Varios"
    ruta_map = {}
    for a in asigs:
        cli = a.item.categoria.proyecto.cliente
        ruta_map[a.item.ruta_item] = (
            a.item.categoria.proyecto.nombre_pmc,
            a.item.nombre,
        )

    for f in req.archivos_incluidos:
        for rbase, (p, i) in ruta_map.items():
            if f.startswith(rbase):
                datos[p][i].append(f)
                break

    # Generar doc dummy
    # (Para no romper lo que ya tenias manual, usamos el generador simple aqui o adaptamos el maestro)
    # Adaptamos maestro rapidito:
    dummy_req = ReporteRequest(
        tipo="Personalizado", fecha_inicio=datetime.now().strftime("%Y-%m-%d")
    )
    dest, fname = _crear_documento_maestro(dummy_req, {cli: datos})

    report_dir = _ensure_reports_dir()

    if not os.path.exists(dest):
        raise HTTPException(500, f"Reporte no generado en disco: {dest}")

    safe_desc = safe_filename(plan.descripcion, default=f"plan_{plan_id}")
    final_name = f"Resumen_{safe_desc}.docx"
    final_path = os.path.join(report_dir, final_name)

    # os.replace es mejor que rename (reemplaza si existe)
    os.replace(dest, final_path)

    # 4. Registrar en Historial (NUEVO para Plan)
    historial = modelos.ReporteHistorial(
        tipo_reporte="plan",
        rango_fechas=datetime.now().strftime("%Y-%m-%d"),
        cliente=cli,
        plan_id=plan_id,
        nombre_archivo=final_name,
        ruta_fisica=final_path,
    )
    db.add(historial)
    db.commit()

    return {"status": "ok", "archivo": final_name}

    bg.add_task(nucleo.run_storage_index_refresh)
    return FileResponse(path=final_path, filename=final_name)


@router.post("/informes/start-resumen-plan/{plan_id}")
def start_informe_custom_plan(
    plan_id: int,
    req: ReportePlanCustomRequest,
    bg: BackgroundTasks,
    db: Session = Depends(dependencias.get_db),
):
    plan = (
        db.query(modelos.PlanTrabajo).filter(modelos.PlanTrabajo.id == plan_id).first()
    )
    if not plan:
        raise HTTPException(404, "Plan no encontrado")

    job_id = reporte_service.create_report_job(db)
    log.info(
        f"[REPORT_START] Job {job_id} creado para plan {plan_id}. Archivos: {len(req.archivos_incluidos)}"
    )

    asig = (
        db.query(modelos.AsignacionPlan)
        .filter(modelos.AsignacionPlan.plan_id == plan_id)
        .first()
    )
    cliente = asig.item.categoria.proyecto.cliente if asig else "Varios"

    bg.add_task(
        _generar_informe_async_worker,
        job_id,
        plan_id,
        req.archivos_incluidos,
        "default",
        plan.descripcion,
        cliente,
    )
    return {"job_id": job_id}


@router.get("/informes/job-status/{job_id}")
def get_report_job_status(job_id: str, db: Session = Depends(dependencias.get_db)):
    job = db.query(modelos.ReportJob).filter(modelos.ReportJob.id == job_id).first()
    if not job:
        raise HTTPException(404, "Job no encontrado")
    return {
        "status": job.status,
        "progress": job.progress,
        "download_url": job.download_url,
        "error": job.error_message,
    }


@router.get("/informes/download-job/{job_id}")
def download_report_job(job_id: str, db: Session = Depends(dependencias.get_db), _g: modelos.User = Depends(dependencias.require_gestion)):
    job = db.query(modelos.ReportJob).filter(modelos.ReportJob.id == job_id).first()
    if not job:
        raise HTTPException(404, "Tarea no encontrada")

    if job.status != "completed":
        raise HTTPException(400, "El informe aún no está listo")

    # Buscar el archivo físico que contenga el job_id
    import os

    try:
        reports_dir = _ensure_reports_dir()
        files = [f for f in os.listdir(reports_dir) if job_id in f]
        if not files:
            log.error(
                f"[DOWNLOAD_ERROR] No se encontró archivo en {reports_dir} para job {job_id}"
            )
            raise HTTPException(404, "El archivo físico no existe en el servidor")

        full_path = os.path.join(reports_dir, files[0])
        return FileResponse(path=full_path, filename=files[0])
    except Exception as e:
        log.error(f"[DOWNLOAD_ERROR] Error al buscar archivo: {e}")
        raise HTTPException(500, "Error interno al recuperar el archivo")


@router.get("/informes/download-direct/{report_id}")
def download_report_direct(report_id: int, db: Session = Depends(dependencias.get_db), _g: modelos.User = Depends(dependencias.require_gestion)):
    reporte = (
        db.query(modelos.ReporteHistorial)
        .filter(modelos.ReporteHistorial.id == report_id)
        .first()
    )
    if not reporte:
        raise HTTPException(404, "Reporte no encontrado en el historial")

    # Path del archivo
    reports_dir = _ensure_reports_dir()
    full_path = os.path.join(reports_dir, reporte.nombre_archivo)
    if not os.path.exists(full_path):
        log.error(f"[DOWNLOAD_ERROR] Archivo no encontrado: {full_path}")
        raise HTTPException(404, "El archivo físico no existe en el servidor")

    return FileResponse(path=full_path, filename=reporte.nombre_archivo)
